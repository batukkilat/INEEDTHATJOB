"""Phase 3: resume tailoring and DOCX generation."""
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sqlmodel import Session

from config import settings
from generation.common import build_profile_json, load_prompt
from utils.llm import chat_with_tool
from utils.logging import get_logger

log = get_logger(__name__)

_TAILOR_TOOL = {
    "description": "Select and tailor resume content for a specific job posting",
    "input_schema": {
        "type": "object",
        "properties": {
            "summary": {"type": "string"},
            "selected_skills": {"type": "array", "items": {"type": "string"}},
            "experiences": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "company":    {"type": "string"},
                        "title":      {"type": "string"},
                        "start_date": {"type": "string"},
                        "end_date":   {"type": "string", "description": "Use 'Present' if current"},
                        "location":   {"type": "string"},
                        "bullets":    {"type": "array", "items": {"type": "string"}},
                    },
                    "required": ["company", "title", "start_date", "end_date", "bullets"],
                },
            },
            "education": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "institution": {"type": "string"},
                        "degree":      {"type": "string"},
                        "field":       {"type": "string"},
                        "start_date":  {"type": "string"},
                        "end_date":    {"type": "string"},
                    },
                    "required": ["institution", "degree", "end_date"],
                },
            },
            "certifications": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "name":   {"type": "string"},
                        "issuer": {"type": "string"},
                        "date":   {"type": "string"},
                    },
                    "required": ["name"],
                },
            },
            "projects": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "name":        {"type": "string"},
                        "description": {"type": "string"},
                        "bullets":     {"type": "array", "items": {"type": "string"}},
                    },
                },
            },
        },
        "required": ["summary", "selected_skills", "experiences", "education"],
    },
}


def tailor_resume(job, session: Session, profile: dict | None = None) -> dict:
    prompt_tmpl = load_prompt("resume_tailor.txt")
    if profile is None:
        profile = build_profile_json(session)
    prompt = (prompt_tmpl
              .replace("{job_title}", job.title)
              .replace("{company}", job.company)
              .replace("{job_description}", (job.description or job.title)[:3000])
              .replace("{profile_json}", json.dumps(profile, ensure_ascii=False)[:6000]))

    log.info("resume_tailor_start", job_id=job.id)
    result = chat_with_tool(
        model=settings.generation_model,
        messages=[{"role": "user", "content": prompt}],
        tool_name="tailor_resume",
        tool_schema=_TAILOR_TOOL,
        max_tokens=3000,
    )
    log.info("resume_tailor_done", job_id=job.id,
             experiences=len(result.get("experiences", [])))
    return result


def _fmt_date(d: str | None) -> str:
    if not d:
        return "Present"
    try:
        return datetime.strptime(d[:7], "%Y-%m").strftime("%b %Y")
    except ValueError:
        return d


def _section_header(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x37, 0x6B, 0xE8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "376BE8")
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx(content: dict, job_id: int) -> str:
    Path(settings.output_dir).mkdir(parents=True, exist_ok=True)
    out_path = str(Path(settings.output_dir) / f"resume_{job_id}.docx")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name_p.add_run(settings.from_name or "Your Name")
    r.bold = True
    r.font.size = Pt(16)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(4)
    contact_p.add_run(settings.from_email or "")

    # Summary
    if content.get("summary"):
        _section_header(doc, "Summary")
        doc.add_paragraph(content["summary"]).paragraph_format.space_after = Pt(2)

    # Skills
    if content.get("selected_skills"):
        _section_header(doc, "Technical Skills")
        doc.add_paragraph(", ".join(content["selected_skills"])).paragraph_format.space_after = Pt(2)

    # Experience
    if content.get("experiences"):
        _section_header(doc, "Experience")
        for exp in content["experiences"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            title_run = p.add_run(f"{exp.get('title', '')} – {exp.get('company', '')}")
            title_run.bold = True
            dates = f"{_fmt_date(exp.get('start_date'))} – {_fmt_date(exp.get('end_date'))}"
            if exp.get("location"):
                dates = f"{exp['location']}  |  {dates}"
            dr = p.add_run(f"\t{dates}")
            dr.font.size = Pt(9)
            dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.2)
                bp.add_run(bullet).font.size = Pt(10)

    # Education
    if content.get("education"):
        _section_header(doc, "Education")
        for edu in content["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            degree = edu.get("degree", "")
            if edu.get("field"):
                degree += f" – {edu['field']}"
            r = p.add_run(f"{degree}  |  {edu.get('institution', '')}")
            r.bold = True
            if edu.get("start_date") or edu.get("end_date"):
                dr = p.add_run(f"\t{_fmt_date(edu.get('start_date'))} – {_fmt_date(edu.get('end_date'))}")
                dr.font.size = Pt(9)

    # Certifications
    if content.get("certifications"):
        _section_header(doc, "Certifications")
        for cert in content["certifications"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.add_run(cert.get("name", "")).bold = True
            meta = " | ".join(filter(None, [cert.get("issuer"), cert.get("date")]))
            if meta:
                p.add_run(f"  ({meta})").font.size = Pt(9)

    # Projects
    if content.get("projects"):
        _section_header(doc, "Projects")
        for proj in content["projects"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.add_run(proj.get("name", "")).bold = True
            if proj.get("description"):
                p.add_run(f" – {proj['description']}")
            for bullet in proj.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.2)
                bp.add_run(bullet).font.size = Pt(10)

    doc.save(out_path)
    log.info("resume_docx_saved", path=out_path)
    return out_path


async def generate_resume(job, session: Session, profile: dict | None = None) -> tuple[str, dict]:
    """Tailor and generate resume DOCX. Returns (docx_path, content_dict)."""
    import asyncio
    content = await asyncio.to_thread(tailor_resume, job, session, profile)
    docx_path = generate_docx(content, job.id)
    return docx_path, content
